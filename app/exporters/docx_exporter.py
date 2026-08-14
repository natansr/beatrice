from pathlib import Path
from docx import Document
from app.exporters.base import BaseExporter


class DocxExporter(BaseExporter):
    format_name, extension = "docx", ".docx"

    def export(self, project, options, destination: Path):
        document = Document()
        if options.include_metadata or options.add_title_page:
            document.add_heading(project.title, 0)
            if project.author: document.add_paragraph(project.author)
            if options.add_title_page: document.add_page_break()
        pages = self.pages(project, options)
        for index, page in enumerate(pages):
            if options.include_page_numbers: document.add_heading(f"Página {page.page_number}", level=2)
            for paragraph in self.text(page, options).split("\n\n"):
                document.add_paragraph(paragraph)
            if options.preserve_page_boundaries and index < len(pages) - 1: document.add_page_break()
        document.save(destination)
        return destination

